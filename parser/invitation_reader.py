from pathlib import Path

from pypdf import PdfReader
from docx import Document


def read_invitation(file_path: str) -> str:
    path = Path(file_path)

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return read_pdf(path)

    if suffix == ".docx":
        return read_docx(path)

    raise ValueError("Chỉ hỗ trợ PDF hoặc DOCX")


def read_pdf(path: Path):

    reader = PdfReader(str(path))

    text = ""

    for page in reader.pages:

        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    return text


def read_docx(path: Path):

    doc = Document(str(path))

    text = []

    for p in doc.paragraphs:

        if p.text.strip():
            text.append(p.text)

    # Table cell text lives in doc.tables, never in doc.paragraphs —
    # confirmed against a real DOCX (742 paragraphs, 27 tables with
    # real data) where table content was silently missing entirely.
    # Appended after all paragraph text (not interleaved back into its
    # original position — acceptable for faithful content extraction,
    # not document-layout reconstruction), each table fenced with a
    # "[Bảng N]" marker so it reads as table data, not prose.
    for table_index, table in enumerate(doc.tables, start=1):

        rows = []

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))

        if rows:
            text.append("\n".join([f"[Bảng {table_index}]", *rows]))

    return "\n".join(text)