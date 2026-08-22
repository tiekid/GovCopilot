"""Tests for parser/invitation_reader.py.

No existing test covered this module before — these are real
assertions against real `pypdf`/`python-docx` behavior (real files
built on disk in a TemporaryDirectory), not mocks, since this is the
one shared file-format-reading module every other component in the
codebase is expected to delegate to.
"""

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from parser.invitation_reader import read_docx, read_invitation


class ReadDocxTests(unittest.TestCase):

    def test_reads_paragraph_text(self) -> None:
        with TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "doc.docx"
            doc = Document()
            doc.add_paragraph("Đoạn văn thứ nhất.")
            doc.add_paragraph("Đoạn văn thứ hai.")
            doc.save(str(path))

            text = read_docx(path)

        self.assertIn("Đoạn văn thứ nhất.", text)
        self.assertIn("Đoạn văn thứ hai.", text)

    def test_reads_table_content_not_just_paragraphs(self) -> None:
        """Regression test for the real bug found 2026-08-22: table cell
        text lives in doc.tables, never in doc.paragraphs — a real DOCX
        with 742 paragraphs and 27 tables had every table's content
        silently missing entirely before this fix."""

        with TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "doc.docx"
            doc = Document()
            doc.add_paragraph("Văn bản trước bảng.")

            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Cơ quan phê duyệt"
            table.cell(0, 1).text = "Sở Xây dựng"
            table.cell(1, 0).text = "Chủ nhiệm đồ án"
            table.cell(1, 1).text = "Nguyễn Văn A"

            doc.save(str(path))

            text = read_docx(path)

        self.assertIn("Văn bản trước bảng.", text)
        self.assertIn("Cơ quan phê duyệt", text)
        self.assertIn("Sở Xây dựng", text)
        self.assertIn("Chủ nhiệm đồ án", text)
        self.assertIn("Nguyễn Văn A", text)

    def test_table_content_is_fenced_and_appended_after_paragraphs(self) -> None:
        with TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "doc.docx"
            doc = Document()
            doc.add_paragraph("Đoạn văn.")
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Nội dung ô bảng"
            doc.save(str(path))

            text = read_docx(path)

        self.assertIn("[Bảng 1]", text)
        self.assertLess(text.index("Đoạn văn."), text.index("[Bảng 1]"))
        self.assertIn("Nội dung ô bảng", text)

    def test_empty_table_contributes_nothing(self) -> None:
        with TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "doc.docx"
            doc = Document()
            doc.add_paragraph("Chỉ có đoạn văn.")
            doc.add_table(rows=2, cols=2)  # every cell left empty
            doc.save(str(path))

            text = read_docx(path)

        self.assertEqual(text, "Chỉ có đoạn văn.")
        self.assertNotIn("[Bảng", text)

    def test_read_invitation_dispatches_docx_and_still_finds_table_content(self) -> None:
        with TemporaryDirectory() as tmp_dir:
            path = Path(tmp_dir) / "doc.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Số liệu trong bảng"
            doc.save(str(path))

            text = read_invitation(str(path))

        self.assertIn("Số liệu trong bảng", text)


if __name__ == "__main__":
    unittest.main()
